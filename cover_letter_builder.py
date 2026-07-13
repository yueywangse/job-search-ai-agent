from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config import COVER_LETTER_DOCX

class CoverLetterBuilder:

    def build(self, resume, cover_letter):
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(45)
        section.right_margin = Pt(45)

        self.add_header(doc, resume)
        self.add_date(doc)
        self.add_greeting(doc)
        self.add_opening(doc, cover_letter)
        self.add_body(doc, cover_letter)
        self.add_closing(doc, cover_letter)
        self.add_signature(doc, resume)

        output = Path(COVER_LETTER_DOCX)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

    def add_header(self, doc, resume):

        p = doc.add_paragraph()

        run = p.add_run(resume.name)
        run.bold = True
        run.font.size = Pt(18)

        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        p.add_run(resume.email)

        p = doc.add_paragraph()
        p.add_run(resume.phone)
        
        github = resume.github.replace("https://", "")
        doc.add_paragraph(f"{resume.linkedin} | {github}")

        p.paragraph_format.space_after = Pt(12)

    def add_date(self, doc):
        from datetime import datetime

        doc.add_paragraph(
            datetime.now().strftime("%B %d, %Y")
        )

        doc.add_paragraph()

    def add_greeting(self, doc):

        p = doc.add_paragraph("Dear Hiring Manager,")
        p.paragraph_format.space_after = Pt(12)

    def add_opening(self, doc, cover_letter):

        p = doc.add_paragraph(cover_letter.opening)
        p.paragraph_format.space_after = Pt(12)

    def add_body(self, doc, cover_letter):

        for paragraph in cover_letter.body:
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(12)

    def add_closing(self, doc, cover_letter):

        p = doc.add_paragraph(cover_letter.closing)
        p.paragraph_format.space_after = Pt(18)

    def add_signature(self, doc, resume):

        doc.add_paragraph("Sincerely,")

        p = doc.add_paragraph()
        run = p.add_run(resume.name)
        run.bold = True