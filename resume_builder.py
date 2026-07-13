from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt
from pathlib import Path

from config import TAILOR_DOCX


class ResumeBuilder:

    def build(self, resume, tailored_resume):
        doc = Document()
        
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(45)
        section.right_margin = Pt(45)

        self.add_header(doc, resume)
        self.add_summary(doc, tailored_resume)
        self.add_skills(doc, tailored_resume)
        self.add_experience(doc, tailored_resume)
        self.add_projects(doc, tailored_resume)
        self.add_education(doc, resume)

        output = Path(TAILOR_DOCX)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

    def add_header(self, doc, resume):

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(resume.name)
        run.bold = True
        run.font.size = Pt(18)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{resume.email} | {resume.phone} | {resume.linkedin} | {resume.github}")
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    def add_section_heading(self, doc, title):

        p = doc.add_paragraph()

        run = p.add_run(title.upper())
        run.bold = True
        run.underline = True
        run.font.size = Pt(13)

        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_summary(self, doc, tailored_resume):

        self.add_section_heading(doc, "Professional Summary")

        p = doc.add_paragraph(tailored_resume.professional_summary)
        p.paragraph_format.space_after = Pt(6)

    def add_skills(self, doc, tailored_resume):

        self.add_section_heading(doc, "Skills")

        skills = " • ".join(tailored_resume.skills)

        p = doc.add_paragraph(skills)
        p.paragraph_format.space_after = Pt(6)

    def add_job(self, doc, job):

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Pt(450), WD_TAB_ALIGNMENT.RIGHT)

        title = p.add_run(job.title)
        title.bold = True

        p.add_run("\t")
        p.add_run(job.dates)

        company = doc.add_paragraph()

        run = company.add_run(f"{job.company} | {job.location}")
        run.italic = True
        
        company.paragraph_format.space_after = Pt(2)

        for bullet in job.bullet_points:

            p = doc.add_paragraph(
                bullet,
                style="List Bullet"
            )

            p.paragraph_format.space_after = Pt(0)

    def add_experience(self, doc, tailored_resume):

        self.add_section_heading(doc, "Experience")

        for job in tailored_resume.work_experience:
            self.add_job(doc, job)

    def add_project(self, doc, project):

        p = doc.add_paragraph()

        run = p.add_run(project.title)
        run.bold = True

        for bullet in project.bullet_points:

            p = doc.add_paragraph(
                bullet,
                style="List Bullet"
            )

            p.paragraph_format.space_after = Pt(0)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_projects(self, doc, tailored_resume):

        self.add_section_heading(doc, "Projects")

        for project in tailored_resume.projects:
            self.add_project(doc, project)

    def add_education(self, doc, resume):

        self.add_section_heading(doc, "Education")

        for edu in resume.education:

            p = doc.add_paragraph()

            degree = p.add_run(edu.degree)
            degree.bold = True

            doc.add_paragraph(f"{edu.university} | {edu.date}")
            
            p.paragraph_format.space_after = Pt(4)