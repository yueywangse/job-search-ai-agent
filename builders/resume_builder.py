from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt

from config import TAILOR_DOCX
from models import Resume, TailoredResume, Experience, Project


class ResumeBuilder:
    """Build and save a tailored resume as a DOCX document."""

    def build(self, resume: Resume, tailored_resume: TailoredResume) -> None:
        """Generate and save the formatted resume document."""

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(45)
        section.right_margin = Pt(45)

        self.add_header(doc, resume)
        self.add_summary(doc, tailored_resume)
        self.add_skills(doc, tailored_resume)
        self.add_experience(doc, tailored_resume)
        self.add_projects(doc, tailored_resume)
        self.add_education(doc, resume)

        output = Path(TAILOR_DOCX)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

    def add_header(self, doc: Document, resume: Resume) -> None:
        """Add the candidate's contact information."""

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name = p.add_run(resume.name)
        name.bold = True
        name.font.size = Pt(18)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = p.add_run(
            f"{resume.email} | {resume.phone} | {resume.linkedin} | {resume.github}"
        )
        contact.font.size = Pt(10)

        p.paragraph_format.space_after = Pt(6)

    def add_section_heading(self, doc: Document, title: str) -> None:
        """Add a section heading."""

        p = doc.add_paragraph()

        heading = p.add_run(title.upper())
        heading.bold = True
        heading.underline = True
        heading.font.size = Pt(13)

        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_summary(self, doc: Document, tailored_resume: TailoredResume) -> None:
        """Add the professional summary."""

        self.add_section_heading(doc, "Professional Summary")

        p = doc.add_paragraph(tailored_resume.professional_summary)
        p.paragraph_format.space_after = Pt(6)

    def add_skills(self, doc: Document, tailored_resume: TailoredResume) -> None:
        """Add the skills section."""

        self.add_section_heading(doc, "Skills")

        skills = " • ".join(tailored_resume.skills)

        p = doc.add_paragraph(skills)
        p.paragraph_format.space_after = Pt(6)

    def add_job(self, doc: Document, job: Experience) -> None:
        """Add a work experience entry."""

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Pt(450), WD_TAB_ALIGNMENT.RIGHT)

        title = p.add_run(job.title)
        title.bold = True

        p.add_run("\t")
        p.add_run(job.dates)

        company = doc.add_paragraph()

        info = company.add_run(f"{job.company} | {job.location}")
        info.italic = True

        company.paragraph_format.space_after = Pt(2)

        for bullet in job.bullet_points:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.paragraph_format.space_after = Pt(0)

    def add_experience(self, doc: Document, tailored_resume: TailoredResume) -> None:
        """Add the work experience section."""

        self.add_section_heading(doc, "Experience")

        for job in tailored_resume.work_experience:
            self.add_job(doc, job)

    def add_project(self, doc: Document, project: Project) -> None:
        """Add a project entry."""

        p = doc.add_paragraph()

        title = p.add_run(project.title)
        title.bold = True

        for bullet in project.bullet_points:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

    def add_projects(self, doc: Document, tailored_resume: TailoredResume) -> None:
        """Add the projects section."""

        self.add_section_heading(doc, "Projects")

        for project in tailored_resume.projects:
            self.add_project(doc, project)

    def add_education(self, doc: Document, resume: Resume) -> None:
        """Add the education section."""

        self.add_section_heading(doc, "Education")

        for education in resume.education:
            p = doc.add_paragraph()

            degree = p.add_run(education.degree)
            degree.bold = True

            p.paragraph_format.space_after = Pt(2)

            p = doc.add_paragraph(f"{education.university} | {education.date}")
            p.paragraph_format.space_after = Pt(4)
