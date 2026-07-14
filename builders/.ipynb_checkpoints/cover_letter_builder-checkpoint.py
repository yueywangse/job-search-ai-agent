from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config import COVER_LETTER_DOCX
from models import CoverLetter, Resume

class CoverLetterBuilder:
    """Build and save a DOCX cover letter from structured cover letter data."""
    
    def build(self, resume: Resume, cover_letter: CoverLetter) -> None:
        """Generate and save the formatted cover letter document."""
        
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(45)
        section.right_margin = Pt(45)

        self.add_header(doc, resume)
        self.add_date(doc)
        self.add_greeting(doc)
        self.add_opening(doc, cover_letter)
        self.add_body(doc, cover_letter)
        self.add_closing(doc, cover_letter)
        self.add_signature(doc, resume)

        output = Path(COVER_LETTER_DOCX)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

    def add_header(self, doc: Document, resume: Resume) -> None:
        """Add the candidate's contact information to the document header."""
        
        p = doc.add_paragraph()

        run = p.add_run(resume.name)
        run.bold = True
        run.font.size = Pt(18)

        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        p.add_run(resume.email)

        p = doc.add_paragraph()
        p.add_run(resume.phone)
        
        github = resume.github.replace("https://", "")
        p = doc.add_paragraph(f"{resume.linkedin} | {github}")

        p.paragraph_format.space_after = Pt(12)

    def add_date(self, doc: Document) -> None:
        """Insert the current date into the cover letter."""
        
        p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        p.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()

    def add_greeting(self, doc: Document) -> None:
        """Add the greeting to the cover letter."""
        
        p = doc.add_paragraph("Dear Hiring Manager,")
        p.paragraph_format.space_after = Pt(12)

    def add_opening(self, doc: Document, cover_letter: CoverLetter) -> None:
        """Add the opening paragraph."""
        
        p = doc.add_paragraph(cover_letter.opening)
        p.paragraph_format.space_after = Pt(12)

    def add_body(self, doc: Document, cover_letter: CoverLetter) -> None:
        """Add the body paragraphs."""
        
        for paragraph in cover_letter.body:
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(12)

    def add_closing(self, doc: Document, cover_letter: CoverLetter) -> None:
        """Add the closing paragraph."""
        
        p = doc.add_paragraph(cover_letter.closing)
        p.paragraph_format.space_after = Pt(18)

    def add_signature(self, doc: Document, resume: Resume) -> None:
        """Add the closing signature with the candidate's name."""
        
        doc.add_paragraph("Sincerely,")

        p = doc.add_paragraph()
        run = p.add_run(resume.name)
        run.bold = True